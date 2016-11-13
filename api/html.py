# -*- coding:utf-8 -*-

import re
import os
import docx
import urllib2
from urllib import urlopen
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from HTMLParser import HTMLParser
from config.config import GLOBAL_CONF

############ html转docx ###############
class HTMLClient:
    #获取html网页源码
    def GetPage(self, url):
        #user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        user_agent = 'Mozilla/5.0 (X11; Linux i686) AppleWebKit/537.36 (KHTML, like Gecko) Ubuntu Chromium/34.0.1847.116 Chrome/34.0.1847.116 Safari/537.36'
        headers = { 'User-Agent' : user_agent }
        req = urllib2.Request(url, None, headers)
        try:
            res = urllib2.urlopen(req)
            return res.read()
        except Exception as e:
            print "GetPage error :", e
            return None

    #获取网络图片并保存在程序运行目录下
    def GetPic(self, url):
        user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        headers = { 'User-Agent' : user_agent }
        req = urllib2.Request(url, None, headers)
        try:
            res = urllib2.urlopen(req)
            return res.read()
        except Exception as e:
            print "GetPic error :", e
            return None

    def GetPicStatic(self, url):
        buf_size = 1024
        image = ''
        urlData = url.split('/')[-1]
        url = '../admin/'+GLOBAL_CONF['ueconfig_dir']+'/'+urlData
        try:
            with open(url, 'r') as f:
                while True:
                    data = f.read(buf_size)
                    if not data:
                        break
                    image += data
        except Exception as e:
            print "GetPicStatic error :", e
            return None
        return image

class MYHTMLParser(HTMLParser):
    def __init__(self, docfile, title):
        HTMLParser.__init__(self)
        self.docfile = docfile
        self.doc = Document()
        self.myclient = HTMLClient()
        self.text = ""
        self.content = ""
        self.validity = False
        self.title = False
        self.isdescription = False
        self.picList=[]
        self.bold=False
        self.p=None

    #根据标签头类型决定标签内容的格式
    def handle_starttag(self, tag, attrs):
        #print "Encountered the beginning of a %s tag" % tag
        self.isdescription = False
        # for (variable, value)  in attrs:
        #     print "tag = {tag}, variable = {var}, value = {str}".format(tag=tag, var=variable, str=value)

        #<h1>标签说明其中的内容是标题
        # title = re.match(r'h(\d)', tag)
        # if title:
        #     self.title = title
        #图片的处理比较复杂，首先需要找到对应的图片的url，然后下载并写入doc中
        #下载的图片格式如果有问题，docx模块会报错，因此重新定义异常处理
        #图片名称需要记录下来，在文档保存后要自动删除
        if tag == "img":
            if len(attrs) == 0: pass
            else:
                for (variable, value)  in attrs:
                    if variable == "src":
                        #此处图片url类型为[http://url/pic.img!200*200]
                        #不同网站图片类型不同，因此当作不同处理
                        url = value.split('!')[0]
                        dir = GLOBAL_CONF['ueconfig_dir']
                        # print "url = ", url, dir
                        if dir in url:
                            picdata = self.myclient.GetPicStatic(url)
                        else:
                            picdata = self.myclient.GetPic(url)
                        if picdata == None:
                            pass
                        else:
                            if len(picdata)<100000:
                                img_width = float('%.4f' % (float(len(picdata))/20480))
                                if img_width < 0.01:
                                    img_width = 0.09
                            else:
                                img_width = 1.5
                            # print "len(picdata) = ",len(picdata)
                            pictmp = value.split('/')[-1].split('!')[0]
                            picfix = value.split('/')[-1].split('!')[-1]
                            with open(pictmp, 'wb') as pic:
                                pic.write(bytes(picdata))
                                pic.close()
                            try:
                                if picfix[0:1] == 'c':
                                    self.doc.add_picture(pictmp, width=Inches(img_width))
                                    # print "add_picture 11= "
                                else:
                                    self.doc.add_picture(pictmp, width=Inches(img_width))
                                    # print "add_picture 22= "
                            except docx.image.exceptions.UnrecognizedImageError as e:
                            # except Exception as e:
                                print "handle_starttag error = ", e
                            self.picList.append(pictmp)
        #javascript脚本
        if tag == 'script':
            self.isdescription = True

        tagList = ['div','br','p','img']
        if tag not in tagList:
            self.validity = True
        else:
            self.validity = False

        # self.doc.add_paragraph('text', RGBColor(0xff, 0x99, 0xcc))
        # print "attrs = ",attrs[0][1] if attrs else'-'
        if( attrs and 'font-weight' in attrs[0][1]) or tag=='strong':
            if self.p == None:
                self.p = self.doc.add_paragraph('')
                self.p.add_run(self.content.decode("utf-8"))
                self.content = ''
            self.bold=True
        else:
            self.bold = False

    def handle_data(self, data):
        # print "data.strip() = ", data.strip()
        if self.title:
            title = self.title.group()
            self.doc.add_heading(data.decode("utf-8"), int(title[1]))
            self.title = False
            return

        if (data and data.strip() != '' and self.bold==False) or self.bold==True:
            self.content += data.strip()
            self.text = ''
            return

    def handle_endtag(self, tag):

        tagList = ['div','br','p','img']
        if (self.content!= '' and self.validity==False and self.bold==False) or (self.content!= '' and tag in tagList and self.bold==False):
            content =self.content.decode("utf-8")
            if self.p:
                # self.p.add_run(content)
                self.p=None
                self.doc.add_paragraph(content)
            else:
                self.doc.add_paragraph(content)
            self.content=''
            # p.add_run(' and some ')

        if self.bold == True and self.p:
            content =self.content.decode("utf-8")
            self.p.add_run(content).bold = True
            self.content=''

    def complete(self, html):
        self.feed(html)
        self.doc.save(self.docfile)
        for item in self.picList:
            if os.path.exists(item):
                os.remove(item)
